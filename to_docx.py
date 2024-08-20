import docx
from docx.shared import Pt
import create_hyperlink
import datetime
import pathlib

def write_articles(selected_articles, document, target = None):
    for a in selected_articles:
        if target is not None:
            try:
                if a.tier_one is None:
                    continue
                if a.tier_one != target:
                    continue
            except:
                continue
        else:
            try:
                if a.tier_one is not None:
                    continue
            except:
                continue
        p = document.add_paragraph("")
        # print(a.Link)
        create_hyperlink.add_hyperlink(p, a.Title, a.Link)
        country = "Germany"
        if a.Source_name == "NZZ":
            country = "Switzerland"
        if a.Date is None:
            a.Date = datetime.datetime.now()
        p.add_run("\n{} ({}), {} {}, {}".format(a.Source_name,country,a.Date.strftime("%B"),a.Date.day,a.Date.year))

        p = document.add_paragraph(a.Leadtext)
        p.style = document.styles['Normal'] 

def set_font(document,font_name,font_size,underline:bool):
    style = document.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.underline = underline

def create_word_document(selected_articles,language):
    document = docx.Document()

    set_font(document,'Arial',11,False)
    p = document.add_paragraph('Media Snapshot Germany')
    p.style = document.styles['Normal']
    p.style = document.styles['Normal']
    p = document.add_paragraph('Find here latest news from Germany.')
    p.style = document.styles['Normal']
    p = document.add_paragraph("")

    mydate = datetime.datetime.now()
    month = mydate.strftime("%B").upper()
    day = mydate.day
    p = document.add_paragraph().add_run('{} {}'.format(month,day))
    p.font.size = Pt(14)

    p = document.add_paragraph().add_run('')
    p.font.size = Pt(11)

    p = document.add_paragraph().add_run('Trending News')
    p.font.size = Pt(14)
    p.font.underline = True

    write_articles(selected_articles, document)

    today = datetime.datetime.now()
    month = today.month
    if month < 10:
        month = '0'+str(month)
    day = today.day
    if day < 10:
        day = '0' + str(day)
    year = today.year - 2000
    if language == "DE":
        document.save(pathlib.Path("Briefings","briefing{}{}{}".format(year,month,day),"briefing{}{}{}_DE.docx".format(year,month,day)))
    elif language == "all":
        document.save(pathlib.Path("Briefings","briefing{}{}{}".format(year,month,day),"briefing{}{}{}_all.docx".format(year,month,day)))
    elif language == "EN":
        document.save(pathlib.Path("Briefings","briefing{}{}{}".format(year,month,day),"briefing{}{}{}.docx".format(year,month,day)))
    else:
        document.save(pathlib.Path("Briefings","briefing{}{}{}".format(year,month,day),"briefing{}{}{}{}.docx".format(year,month,day,language)))